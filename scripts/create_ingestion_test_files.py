from __future__ import annotations

import base64
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


DESKTOP = Path.home() / "Desktop"
OUTPUT = DESKTOP / "\u5bfc\u5165\u6d4b\u8bd5\u6587\u4ef6"
LONG_PARAGRAPH = ("\u672c\u6d4b\u8bd5\u6bb5\u843d\u7528\u4e8e\u9a8c\u8bc1\u4e2d\u6587\u5206\u5757\u903b\u8f91\u3002" * 200) + "\u7ed3\u675f\u3002"


def write_markdown_files() -> None:
    (OUTPUT / "01-normal-markdown.md").write_text(
        "# \u6d4b\u8bd5\u77e5\u8bc6\u6587\u6863\n\n## \u5b66\u4e60\u4e3b\u9898\n\n"
        "\u672c\u6587\u6863\u7528\u4e8e\u9a8c\u8bc1 Markdown \u6587\u4ef6\u7684\u6b63\u5e38\u5bfc\u5165\u3001\u6807\u9898\u8def\u5f84\u4fdd\u7559\u548c\u5411\u91cf\u7d22\u5f15\u3002\n\n"
        "## \u5173\u952e\u95ee\u9898\n\n\u5b66\u751f\u5982\u4f55\u5c06\u5f53\u5730\u7ea2\u8272\u8d44\u6e90\u4e0e\u5386\u53f2\u5b66\u4e60\u8fde\u63a5\u8d77\u6765\uff1f\n",
        encoding="utf-8",
    )
    (OUTPUT / "02-long-chinese-no-whitespace.md").write_text(
        "# \u4e2d\u6587\u5206\u5757\u538b\u6d4b\n\n## \u8d85\u957f\u6bb5\u843d\n\n" + LONG_PARAGRAPH + "\n",
        encoding="utf-8",
    )
    (OUTPUT / "04-empty-text-expected-failure.md").write_text("", encoding="utf-8")


def write_test_image() -> Path:
    image = OUTPUT / "test-image.png"
    # A compact valid 1x1 PNG is sufficient to exercise DOCX media extraction.
    image.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQIHWP4z8DwHwAFgAI/ScL3tgAAAABJRU5ErkJggg=="
    ))
    return image


def write_docx(image: Path) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    document.add_heading("\u542b\u56fe\u7247\u7684 DOCX \u5bfc\u5165\u6d4b\u8bd5", level=0)
    document.add_heading("\u6587\u672c\u5185\u5bb9", level=1)
    document.add_paragraph("\u8fd9\u4e2a\u6587\u6863\u7528\u4e8e\u9a8c\u8bc1 DOCX \u8f6c Markdown\u3001\u6807\u9898\u63d0\u53d6\u548c\u5185\u5d4c\u56fe\u7247\u5904\u7406\u3002")
    document.add_heading("\u56fe\u7247\u6d4b\u8bd5", level=1)
    document.add_picture(str(image), width=Inches(0.25))
    document.add_paragraph("\u56fe\u7247\u4e0d\u9700\u8981\u89c6\u89c9\u6a21\u578b\u624d\u80fd\u4f7f\u6587\u672c\u5bfc\u5165\u6210\u529f\uff1b\u6a21\u578b\u4e0d\u53ef\u7528\u65f6\u9884\u671f\u7ed3\u679c\u4e3a DEGRADED\u3002")
    document.save(OUTPUT / "03-docx-with-image.docx")


def write_pdf() -> None:
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    pdfmetrics.registerFont(TTFont("MicrosoftYaHei", str(font_path), subfontIndex=0))
    pdf = canvas.Canvas(str(OUTPUT / "05-pdf-native-fallback-test.pdf"), pagesize=A4)
    pdf.setFont("MicrosoftYaHei", 16)
    pdf.drawString(72, 790, "PDF \u672c\u5730\u89e3\u6790\u964d\u7ea7\u6d4b\u8bd5")
    pdf.setFont("MicrosoftYaHei", 11)
    lines = [
        "\u6b64 PDF \u7528\u4e8e\u9a8c\u8bc1 MinerU \u672a\u914d\u7f6e\u6216\u5931\u8d25\u65f6\uff0c\u7cfb\u7edf\u80fd\u56de\u9000\u5230\u672c\u5730 PDF \u6587\u672c\u63d0\u53d6\u3002",
        "\u9884\u671f\uff1a\u6587\u6863\u6700\u7ec8\u4e3a DEGRADED\uff0c\u800c\u4e0d\u662f FAILED\uff0c\u4e14\u80fd\u4ea7\u751f\u975e\u96f6\u6570\u91cf\u7684\u5206\u5757\u4e0e\u5411\u91cf\u3002",
    ]
    y = 750
    for line in lines:
        pdf.drawString(72, y, line)
        y -= 32
    pdf.save()


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    write_markdown_files()
    image = write_test_image()
    write_docx(image)
    write_pdf()
    print(OUTPUT)


if __name__ == "__main__":
    main()
